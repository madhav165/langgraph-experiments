import os
import getpass
import sqlite3
from typing import Annotated
from langchain_openai import ChatOpenAI
from langchain_community.tools import DuckDuckGoSearchRun
from langchain_core.messages import BaseMessage, AIMessage
from typing_extensions import TypedDict
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.sqlite import SqliteSaver
from langchain.tools import tool
from datetime import datetime
from docx import Document
from PIL import Image
import io

from dotenv import load_dotenv

load_dotenv("../../.venv")

# def _set_env(var: str):
#     if not os.environ.get(var):
#         os.environ[var] = getpass.getpass(f"{var}: ")

# os.environ["LANGCHAIN_TRACING_V2"] = True
# os.environ["LANGCHAIN_ENDPOINT"] = "https://api.smith.langchain.com"
# os.environ["LANGCHAIN_PROJECT"] = "LangGraph Tutorial"

config = {"configurable": {"thread_id": "1"}}
memory = SqliteSaver.from_conn_string("test.sqlite")

class State(TypedDict):
    messages: Annotated[list, add_messages]

graph_builder = StateGraph(State)

@tool
def get_current_date() -> str:
    """Return the current date in YYYY-MM-DD format."""
    return str(datetime.today().date())

def open_or_create_doc(file_path):
    """
    Open a Word document if it exists; otherwise, create a new one.
    """
    if os.path.exists(file_path):
        # Open the existing document
        doc = Document(file_path)
        print(f"Opened existing document: {file_path}")
    else:
        # Create a new document
        doc = Document()
        doc.add_heading('New Document', 0)
        doc.save(file_path)
        print(f"Created new document: {file_path}")
    
    return doc

@tool
def write_to_docx(file_name: str, input: str) -> None:
    """Append given input to file_name"""
    doc = open_or_create_doc(file_name)
    doc.add_paragraph(input)
    doc.save(file_name)
    # return str(datetime.today().date())

ddg_tool = DuckDuckGoSearchRun(max_results=2)
tools = [get_current_date, write_to_docx, ddg_tool]
llm = ChatOpenAI(model="gpt-3.5-turbo")
# llm = ChatOpenAI(model="gpt-4o")
llm_with_tools = llm.bind_tools(tools)

def chatbot(state: State):
    return {"messages": [llm_with_tools.invoke(state["messages"])]}
graph_builder.add_node("chatbot", chatbot)

tool_node = ToolNode(tools=[get_current_date, write_to_docx, ddg_tool])
graph_builder.add_node("tools", tool_node)

graph_builder.add_conditional_edges("chatbot", tools_condition)
graph_builder.add_edge("tools", "chatbot")
graph_builder.add_edge(START, "chatbot")

graph = graph_builder.compile(
    checkpointer=memory,
)

# try:
#     # Assuming graph.get_graph().draw_mermaid_png() returns image data in bytes
#     image_data = graph.get_graph().draw_mermaid_png()

#     # Convert bytes data to an image object
#     image = Image.open(io.BytesIO(image_data))

#     # Save the image to disk
#     image.save('graph.png')
#     print("Image saved to graph.png")
# except Exception as e:
#     # This requires some extra dependencies and is optional
#     print(f"An error occurred: {e}")
#     pass

# snapshot = graph.get_state(config)
# print(snapshot)

# exit()



while True:
    user_input = input("User: ")
    if user_input.lower() in ["quit", "exit", "q"]:
        print("Goodbye!")
        break
    for event in graph.stream({"messages": [("user", user_input)]}, config):
        for value in event.values():
            if isinstance(value["messages"][-1], BaseMessage):
                print("Assistant:", value["messages"][-1].content)
                # snapshot = graph.get_state(config)
                # existing_message = snapshot.values["messages"][-1]
                if isinstance(value["messages"][-1], AIMessage):
                    print(value["messages"][-1].tool_calls)
                    # graph.stream(None, config, stream_mode="values")

# snapshot = graph.get_state(config)
# print(snapshot)