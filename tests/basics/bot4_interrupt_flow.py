#%%
import os
import sqlite3
from typing import Annotated
from langchain_openai import ChatOpenAI
from langchain_ollama import ChatOllama
from langchain_experimental.llms.ollama_functions import OllamaFunctions
from langchain_community.tools import DuckDuckGoSearchRun
from langchain_core.messages import BaseMessage, AIMessage
from typing_extensions import TypedDict
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.sqlite import SqliteSaver
from langchain.tools import tool
from datetime import datetime
from docx import Document
from PIL import Image
import io

from dotenv import load_dotenv

load_dotenv("../../.env")

from llm import *

#%%
config = {"configurable": {"thread_id": "1"}}
memory = SqliteSaver.from_conn_string("test.sqlite")

#%%
class State(TypedDict):
    messages: Annotated[list, add_messages]

#%%
graph_builder = StateGraph(State)

#%%
@tool
def get_current_date() -> str:
    """Return the current date in YYYY-MM-DD format."""
    return str(datetime.today().date())

#%%
def open_or_create_doc(file_path):
    """
    Open a Word document if it exists; otherwise, create a new one.
    """
    if os.path.exists(file_path):
        # Open the existing document
        doc = Document(file_path)
        print(f"Opened existing document: {file_path}")
    else:
        # Create a new document
        doc = Document()
        doc.add_heading('New Document', 0)
        doc.save(file_path)
        print(f"Created new document: {file_path}")
    
    return doc

#%%
@tool
def write_to_docx(file_name: str, input: str) -> None:
    """Append given input to file_name"""
    doc = open_or_create_doc(file_name)
    doc.add_paragraph(input)
    doc.save(file_name)
    # return str(datetime.today().date())

#%%
ddg_tool = DuckDuckGoSearchRun(max_results=2)
tools = [get_current_date, write_to_docx, ddg_tool]
llm = nvidia_llm
llm_with_tools = llm.bind_tools(tools)

#%%
def chatbot(state: State):
    return {"messages": [llm_with_tools.invoke(state["messages"])]}
graph_builder.add_node("chatbot", chatbot)

#%%
tool_node = ToolNode(tools=[get_current_date, write_to_docx, ddg_tool])
graph_builder.add_node("tools", tool_node)

#%%
graph_builder.add_conditional_edges("chatbot", tools_condition)
graph_builder.add_edge("tools", "chatbot")
graph_builder.add_edge(START, "chatbot")

#%%
graph = graph_builder.compile(
    checkpointer=memory,
    interrupt_before=["tools"],
)

# try:
#     # Assuming graph.get_graph().draw_mermaid_png() returns image data in bytes
#     image_data = graph.get_graph().draw_mermaid_png()

#     # Convert bytes data to an image object
#     image = Image.open(io.BytesIO(image_data))

#     # Save the image to disk
#     image.save('graph.png')
#     print("Image saved to graph.png")
# except Exception as e:
#     # This requires some extra dependencies and is optional
#     print(f"An error occurred: {e}")
#     pass

# snapshot = graph.get_state(config)
# print(snapshot)

# exit()

#%%
user_input = "I'm learning LangGraph. Could you do some research on it for me?"
config = {"configurable": {"thread_id": "1"}}

#%%
# The config is the **second positional argument** to stream() or invoke()!
events = graph.stream(
    {"messages": [("user", user_input)]}, config, stream_mode="values"
)
for event in events:
    if "messages" in event:
        event["messages"][-1].pretty_print()

#%%
snapshot = graph.get_state(config)
print(snapshot.next)


#%%
existing_message = snapshot.values["messages"][-1]
print(existing_message.tool_calls)
# existing_message.pretty_print()


#%%
snapshot = graph.get_state(config)
existing_message = snapshot.values["messages"][-1]
print("Original")
print("Message ID", existing_message.id)
print(existing_message.tool_calls[0])
new_tool_call = existing_message.tool_calls[0].copy()
new_tool_call["args"]["query"] = "LangGraph human-in-the-loop workflow"
new_message = AIMessage(
    content=existing_message.content,
    tool_calls=[new_tool_call],
    # Important! The ID is how LangGraph knows to REPLACE the message in the state rather than APPEND this messages
    id=existing_message.id,
)

print("Updated")
print(new_message.tool_calls[0])
print("Message ID", new_message.id)
graph.update_state(config, {"messages": [new_message]})

print("\n\nTool calls")
graph.get_state(config).values["messages"][-1].tool_calls




#%%
events = graph.stream(None, config, stream_mode="values")
for event in events:
    if "messages" in event:
        event["messages"][-1].pretty_print()


# %%
events = graph.stream(
    {
        "messages": (
            "user",
            "Remember what I'm learning about?",
        )
    },
    config,
    stream_mode="values",
)
for event in events:
    if "messages" in event:
        event["messages"][-1].pretty_print()
# %%
