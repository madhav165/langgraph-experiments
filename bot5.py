#%%
import os
import sqlite3
from typing import Annotated
from langchain_openai import ChatOpenAI
from langchain_ollama import ChatOllama
from langchain_experimental.llms.ollama_functions import OllamaFunctions
from langchain_community.tools import DuckDuckGoSearchRun
from langchain_core.pydantic_v1 import BaseModel
from langchain_core.messages import BaseMessage, AIMessage, ToolMessage
from typing_extensions import TypedDict
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.sqlite import SqliteSaver
from langchain.tools import tool
from datetime import datetime
from docx import Document
from PIL import Image
import io

#%%
config = {"configurable": {"thread_id": "1"}}
memory = SqliteSaver.from_conn_string("test.sqlite")

#%%
class State(TypedDict):
    messages: Annotated[list, add_messages]
    ask_human: bool

#%%
class RequestAssistance(BaseModel):
    """Escalate the conversation to an expert. Use this if you are unable to assist directly or if the user requires support beyond your permissions.

    To use this function, relay the user's 'request' so the expert can provide the right guidance.
    """

    request: str

#%%
graph_builder = StateGraph(State)

#%%
@tool
def get_current_date() -> str:
    """Return the current date in YYYY-MM-DD format."""
    return str(datetime.today().date())

#%%
def open_or_create_doc(file_path):
    """
    Open a Word document if it exists; otherwise, create a new one.
    """
    if os.path.exists(file_path):
        # Open the existing document
        doc = Document(file_path)
        print(f"Opened existing document: {file_path}")
    else:
        # Create a new document
        doc = Document()
        doc.add_heading('New Document', 0)
        doc.save(file_path)
        print(f"Created new document: {file_path}")
    
    return doc

#%%
@tool
def write_to_docx(file_name: str, input: str) -> None:
    """Append given input to file_name"""
    doc = open_or_create_doc(file_name)
    doc.add_paragraph(input)
    doc.save(file_name)
    # return str(datetime.today().date())

#%%
ddg_tool = DuckDuckGoSearchRun(max_results=2)
tools = [get_current_date, write_to_docx, ddg_tool]
llm = ChatOpenAI(model="gpt-3.5-turbo")
# llm = ChatOllama(model="phi3")
llm_with_tools = llm.bind_tools(tools + [RequestAssistance])

#%%
def chatbot(state: State):
    response = llm_with_tools.invoke(state["messages"])
    ask_human = False
    if (
        response.tool_calls
        and response.tool_calls[0]["name"] == RequestAssistance.__name__
    ):
        ask_human = True
    return {"messages": [response], "ask_human": ask_human}

graph_builder.add_node("chatbot", chatbot)

#%%
def create_response(response: str, ai_message: AIMessage):
    return ToolMessage(
        content=response,
        tool_call_id=ai_message.tool_calls[0]["id"],
    )

#%%
def human_node(state: State):
    new_messages = []
    if not isinstance(state["messages"][-1], ToolMessage):
        # Typically, the user will have updated the state during the interrupt.
        # If they choose not to, we will include a placeholder ToolMessage to
        # let the LLM continue.
        new_messages.append(
            create_response("No response from human.", state["messages"][-1])
        )
    return {
        # Append the new messages
        "messages": new_messages,
        # Unset the flag
        "ask_human": False,
    }

graph_builder.add_node("human", human_node)

#%%
def select_next_node(state: State):
    if state["ask_human"]:
        return "human"
    # Otherwise, we can route as before
    return tools_condition(state)

graph_builder.add_conditional_edges(
    "chatbot",
    select_next_node,
    {"human": "human", "tools": "tools", "__end__": "__end__"},
)



#%%
tool_node = ToolNode(tools=[get_current_date, write_to_docx, ddg_tool])

graph_builder.add_node("tools", tool_node)

#%%
graph_builder.add_edge("tools", "chatbot")
graph_builder.add_edge("human", "chatbot")
graph_builder.add_edge(START, "chatbot")

#%%
graph = graph_builder.compile(
    checkpointer=memory,
    interrupt_before=["human"],
)

#%%
try:
    # Assuming graph.get_graph().draw_mermaid_png() returns image data in bytes
    image_data = graph.get_graph().draw_mermaid_png()

    # Convert bytes data to an image object
    image = Image.open(io.BytesIO(image_data))

    # Save the image to disk
    image.save('graph.png')
    print("Image saved to graph.png")
except Exception as e:
    # This requires some extra dependencies and is optional
    print(f"An error occurred: {e}")
    pass

#%%
# snapshot = graph.get_state(config)
# print(snapshot)

# exit()

#%%
# user_input = "I'm learning LangGraph. Could you do some research on it for me?"
user_input = "I need some expert guidance for building this AI agent. Could you request assistance for me?"
config = {"configurable": {"thread_id": "1"}}

#%%
# The config is the **second positional argument** to stream() or invoke()!
events = graph.stream(
    {"messages": [("user", user_input)]}, config, stream_mode="values"
)
for event in events:
    if "messages" in event:
        event["messages"][-1].pretty_print()

#%%
snapshot = graph.get_state(config)
print(snapshot.next)


#%%
existing_message = snapshot.values["messages"][-1]
print(existing_message.tool_calls)
# existing_message.pretty_print()

#%%
ai_message = snapshot.values["messages"][-1]
human_response = (
    "We, the experts are here to help! We'd recommend you check out LangGraph to build your agent."
    " It's much more reliable and extensible than simple autonomous agents."
)
tool_message = create_response(human_response, ai_message)
graph.update_state(config, {"messages": [tool_message]})


#%%
graph.get_state(config).values["messages"]


#%%
events = graph.stream(None, config, stream_mode="values")

for event in events:
    if "messages" in event:
        event["messages"][-1].pretty_print()
